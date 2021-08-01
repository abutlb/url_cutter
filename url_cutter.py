import os
import subprocess
import sys
import time
import webbrowser

try:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.table import WD_TABLE_DIRECTION
    import pandas as pd
    import qrcode
    import requests, json
    from mechanize import Browser
except:
    looping = True
    while looping:
        a = input("""sorry you don't have the modules needed for this program
    pandas, qrcode, requests, json, mechanize, docx
    do you want to install them? (Y/N) >>>""")
        if a.lower() in "y":
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-r", "requirements.txt"])
            looping = False
        elif a.lower() in "n":
            print("i can't run the program without them...sorry ^_^")
            time.sleep(2)
            quit()
        else:
            print("not recognized!! try again")

br = Browser()
br.addheaders = [('User-agent', 'Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.9.0.1) Gecko/2008071615 Fedora/3.0.1-1.fc9 Firefox/3.0.1')]
br.set_handle_robots(False)
br.set_handle_equiv(False)
begin_msg = f"""******************************
made by @abutlb
this shortening url is using cutt.us api 
1.start
2.script page
3.quit
******************************"""
print(begin_msg)

def make_short(url):
    if not os.path.exists("out"):
        os.makedirs("out")
    i = 0
    url['shorted url'] = {}
    url['qrcode'] = {}
    url['title'] = {}
    word_doc = Document()
    word_table = word_doc.add_table(0, 0)
    word_table.direction = WD_TABLE_DIRECTION.RTL
    for index in url.keys():
        word_table.add_column(Cm(5))
        if i == 0:
            word_table.add_row()
        row = word_table.rows[0]
        row.cells[i].text = str(index)
        i += 1
    i = 0
    for url_num, url_link in url['url'].items():
        print(f"working on {url_num + 1} of {len(url['url'].items())}")
        link = requests.get(f'https://cutt.us/api.php?url={url_link}&format=json')
        link2 = json.loads(json.dumps(link.json()))
        url['shorted url'][i] = link2['shorturl']
        qr_code = qrcode.make(url['shorted url'][i])
        br.open(url['url'][i])
        t = br.title()
        qr_code.save(f'out\{t}.jpg')
        url['title'][i] = t
        word_table.add_row()
        row = word_table.rows[i + 1]
        row.cells[0].text = str(url['url'][i])
        row.cells[1].text = str(url['shorted url'][i])
        paragraph = row.cells[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(f'out\{t}.jpg', width=Cm(2), height=Cm(2))
        row.cells[3].text = str(url['title'][i])
        i += 1
        print("done!!")
    df = pd.DataFrame(url)
    df.to_excel('out\out.xlsx', sheet_name='out')
    word_doc.save('out\out.docx')


while True:
    option = input(">>>>").lower()
    if option in "1.start":
        filename = input("ok ...give me the filename (only txt and xlsx are supported right now!!): ")
        if filename.endswith("txt"):
            txt_file = list(open(filename, "r").read().split(","))
            url = {"url":{}}
            for i in range(len(txt_file)):
                url['url'][i] = txt_file[i]
            make_short(url)
        elif filename.endswith("xlsx"):
            sheet_num = input("give me the sheet number :")
            xlsx_file = pd.ExcelFile(filename)
            url = xlsx_file.parse(xlsx_file.sheet_names[int(sheet_num)]).to_dict()
            make_short(url)
        else:
            print("file type not supported!!!")


    elif option in "2":
        webbrowser.open("", new=0)

    elif option in "3.quit":
        for l in range(3,0,-1):
            print(f"quitting in {l}", end="\r")
            time.sleep(1)
        print("bye ^_^                    ")
        time.sleep(1)
        quit()
    else:
        print("not recognized!!!")
